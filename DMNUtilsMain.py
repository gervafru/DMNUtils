from operator import truediv
import uuid
import xml.etree.ElementTree as ET
import lxml.etree as ETL
import openpyxl
from openpyxl.styles import Border, Side, Alignment, Font, PatternFill
from PyQt5 import QtWidgets, uic, QtGui
from PyQt5.QtWidgets import QApplication, QWidget, QInputDialog, QLineEdit, QFileDialog, QMessageBox
from PyQt5.QtGui import *
from PyQt5.QtCore import *
import sys
import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class Ui(QtWidgets.QDialog):

    def __init__(self, *args, **kwargs):
        super(Ui, self).__init__(*args, **kwargs) # Call the inherited classes __init__ method
        uic.loadUi('GUI\\main.ui', self) # Cargar el archivo .ui de PyQT5
        self.setWindowIcon(QtGui.QIcon('GUI\\icon.png'))

     # Conectar los botones con sus respectivas funciones
        self.BotonSalir.clicked.connect(self.Salir)
        self.BotonAbrirDMN.clicked.connect(self.AbrirDMN)
        self.BotonAbrirExcel.clicked.connect(self.AbrirExcel)
        self.BotonReemplazar.clicked.connect(self.Procesar)
        self.BotonExportar.clicked.connect(self.ExportarExcel)
        self.BotonExplicarDMN.clicked.connect(self.ExplicarDMN)

        # Iniciar Gui (ventana principal)
        self.show()


    def Salir(self):
        app.exit
        quit()
        
    def AbrirDMN(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self,"QFileDialog.getOpenFileName()", "","Archivos Excel (*.dmn)", options=options)
        self.ListaTablas.clear()
        self.TextoArchivoDMN.clear()
        self.BotonExportar.setEnabled(False)
        self.BotonExplicarDMN.setEnabled(False)
        if fileName:

            # Controlar que sea un DMN valido y armar lista de tablas con sus ID
            mytree = ET.parse(fileName)
            myroot = mytree.getroot()
            TagGenerico = '{http://www.omg.org/spec/DMN/20180521/MODEL/}'

            global ListaTablasDecision

            ListaTablasDecision = []

            RutaBuscarTablas = './/*' + TagGenerico + 'decisionTable'
            IdTabla = ''

            for Tabla in myroot.findall(RutaBuscarTablas): # Loop por cada decisionTable que encuentra
            #_____________________________________________________________________________________________
                
                IdTabla = Tabla.attrib['id']
                RutaNombreTabla = './/*[@id="' + IdTabla + '"]/..' # Buscar nombre en tablas puras
                
                TipoTabla = ''
                Variable = ''

                BuscarNombre = myroot.find(RutaNombreTabla)

                try:
                    NombreTabla = BuscarNombre.attrib['name'] # Guardar nombre en tablas puras
                    TipoTabla = 'decision'
                except:
                    RutaNombreContextTabla = './/*[@id="' + IdTabla + '"]/../../..' # Buscar nombre en tablas dentro de context
                    BuscarNombreContext = myroot.find(RutaNombreContextTabla)
                    try:
                        NombreTabla = BuscarNombreContext.attrib['name']  # Guardar nombre en tablas dentro de context
                        TipoTabla = 'context'
                        RutaNombreVariableTabla = './/*[@id="' + IdTabla + '"]/../' + TagGenerico + 'variable'
                        BuscarNombreVariable = myroot.find(RutaNombreVariableTabla)
                        Variable = BuscarNombreVariable.attrib['name']
                        NombreTabla = NombreTabla + '/' + Variable
                    except:
                        RutaNombreFuncionTabla = './/*[@id="' + IdTabla + '"]/../../../..' # Buscar nombre tablas dentro de Function
                        BuscarNombreFuncion = myroot.find(RutaNombreFuncionTabla)
                        NombreTabla = BuscarNombreFuncion.attrib['name']
                        TipoTabla = 'function'
                        RutaNombreVariableTabla = './/*[@id="' + IdTabla + '"]/../' + TagGenerico + 'variable'
                        BuscarNombreVariable = myroot.find(RutaNombreVariableTabla)
                        Variable = BuscarNombreVariable.attrib['name']
                        NombreTabla = NombreTabla + '/' + Variable


                # Reconstruir ruta de acceso completa a la decisionTable

                RutaTabla0 = './/*[@id="' + IdTabla + '"]/'
                
                NodosSinId = ['contextEntry']

                RutaTabla0 = './/*[@id="' + IdTabla + '"]'

                Buscar0 = myroot.find(RutaTabla0)

                LRutaAccesoTabla = []

                ParteRuta = str('/' + Buscar0.tag + "[@id='" + Buscar0.attrib['id'] + "']")
                LRutaAccesoTabla.append(ParteRuta)

                ParteRuta = ''
                BuscarAnterior = '/..'

                while Buscar0 != None:
                    
                    Buscar0 = myroot.find(RutaTabla0 + BuscarAnterior)
                    
                    if Buscar0 != None:
                        
                        if TagGenerico in str(Buscar0.tag):
                            NodoEncontrado = str(Buscar0.tag).replace(TagGenerico,'')
                        else:
                            NodoEncontrado = str(Buscar0.tag)

                        if NodoEncontrado not in NodosSinId:
                            ParteRuta = str('/' + Buscar0.tag + "[@id='" + Buscar0.attrib['id'] + "']")
                            LRutaAccesoTabla.append(ParteRuta)
                        else:
                            ParteRuta = str('/' + Buscar0.tag)
                            LRutaAccesoTabla.append(ParteRuta)

                        BuscarAnterior = BuscarAnterior + '/..'

                RutaAcceso = '.'

                for I, PRuta in enumerate(LRutaAccesoTabla):
                    if 'definitions' not in LRutaAccesoTabla[len(LRutaAccesoTabla)-(I+1)]: # Evitar nodo inicial definitions
                        RutaAcceso = RutaAcceso + LRutaAccesoTabla[len(LRutaAccesoTabla)-(I+1)]

                ListaParcial = []
                ListaParcial.append(NombreTabla)
                ListaParcial.append(IdTabla)
                ListaParcial.append(RutaAcceso)
                ListaParcial.append(TipoTabla)
                ListaParcial.append(Variable)
                ListaTablasDecision.append(ListaParcial)

            if ListaTablasDecision == []:
                self.MensajeError = QMessageBox()
                self.MensajeError.setIcon(QMessageBox.Critical)
                self.MensajeError.setWindowTitle("Error")
                self.MensajeError.setText("DMN incorrecto o sin tabla de decisión.")
                self.MensajeError.show()
            else:
                self.TextoArchivoDMN.setText(fileName)
                for nomTabla in ListaTablasDecision:
                    self.ListaTablas.addItem(nomTabla[0])
                self.BotonExportar.setEnabled(True)
                self.BotonExplicarDMN.setEnabled(True)


    def AbrirExcel(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self,"QFileDialog.getOpenFileName()", "","Archivos Excel (*.xlsx)", options=options)
        if fileName:
            self.TextoArchivoExcel.setText(fileName)

    

    def Procesar(self):

# Leer archivo Excel con tabla de decisiones a importar
        def LeerTablaExcel(NombreArchivo):

            ListaResultado = []
            ListaInput = []
            ListaOutput = []
            ListaRules = []

            FormatoCorrecto = True

            ArchivoXLSX = openpyxl.load_workbook(NombreArchivo)
            sheet = ArchivoXLSX.active

            CantColumnas = sheet.max_column
            CantLineas = sheet.max_row

            for Linea in sheet.iter_rows(max_row=3):
                for Celda in Linea:
                    if Celda.row == 1: # Linea 1 = input/output
                        if Celda.value == 'input':
                            ListaTmp = []
                            ListaTmp.append(Celda.value)
                            ListaInput.append(ListaTmp)
                        elif Celda.value == 'output':
                            ListaTmp = []
                            ListaTmp.append(Celda.value)
                            ListaOutput.append(ListaTmp)
                        elif Celda.value not in ['input', 'output']:
                            FormatoCorrecto = False

                    if Celda.row == 2: # Linea 1 = tipo variable (string, number, boolean, any)
                        if Celda.value in ['string', 'number', 'boolean', 'any']:
                            if Celda.column <= len(ListaInput):
                                ListaInput[Celda.column-1].append(Celda.value)
                            elif Celda.column > len(ListaInput) and Celda.column <= (len(ListaInput)+len(ListaOutput)):
                                ListaOutput[((Celda.column)-len(ListaInput))-1].append(Celda.value)
                            elif Celda.column > (len(ListaInput)+len(ListaOutput)):
                                FormatoCorrecto = False
                        else:
                            FormatoCorrecto = False


                    if Celda.row == 3: # Nombre de variables
                        if Celda.value not in ['', None]:
                            if Celda.column <= len(ListaInput):
                                ListaInput[Celda.column-1].append(Celda.value)
                            elif Celda.column > len(ListaInput) and Celda.column <= (len(ListaInput)+len(ListaOutput)):
                                ListaOutput[((Celda.column)-len(ListaInput))-1].append(Celda.value)
                            elif Celda.column > (len(ListaInput)+len(ListaOutput)):
                                FormatoCorrecto = False
                        else:
                            FormatoCorrecto = False


            for Linea in sheet.iter_rows(min_row=4): # Linea 4 en adelante = reglas
                ListaTmp = []
                for Celda in Linea:
                    if Celda.value not in ['', None]:
                        ListaTmp.append(str(Celda.value))
                    else:
                        FormatoCorrecto = False

                ListaRules.append(ListaTmp)


            if FormatoCorrecto == True:
                ListaResultado.append(ListaInput)
                ListaResultado.append(ListaOutput)
                ListaResultado.append(ListaRules)
            else:
                ListaResultado = []
                    

            return ListaResultado
            

        # Eliminar dentro de la decision elegida todos los datos del tipo elegido (input, output, etc.)
        def EliminarSeccion(RutaTabla, TagGenerico, TipoSeccion):        
            
            if TipoSeccion in ['input', 'output', 'rule']:
                for Hijo in myroot.findall(RutaTabla + '/' + TagGenerico + TipoSeccion):
                    Padre = myroot.find(RutaTabla + '/' + TagGenerico + TipoSeccion + "[@id='" + Hijo.attrib['id'] + "']/..")
                    Padre.remove(Hijo)
            
            elif TipoSeccion == 'annotation':
                for Hijo in myroot.findall(RutaTabla + '/' + TagGenerico + TipoSeccion):
                    Padre = myroot.find(RutaTabla + '/' + TagGenerico + TipoSeccion + "[@name='" + Hijo.attrib['name'] + "']/..")
                    Padre.remove(Hijo)

            return


        # Funcion para agregar nuevos Input
        def AgregarInput(RutaAcceso, NombreVariable, TipoVariable):
                
            # Ejemplo input original = '<ns0:input xmlns:ns0="http://www.omg.org/spec/DMN/20180521/MODEL/" id="_87040B53-1DC4-4644-AD36-FA0D7E13D1C8">\n        <ns0:inputExpression id="_BBF28E82-6D8F-458D-A5CD-522DECEFB16F" typeRef="string">\n          <ns0:text>Asociado.car</ns0:text>\n        </ns0:inputExpression>\n      </ns0:input>\n      '
            NivelTabla = myroot.find(RutaAcceso)
            NuevoElemento = ET.Element('ns0:input') # Nivel 1 input
            Id_Random = '_' + str(uuid.uuid4()).upper()
            NuevoElemento.set('id', Id_Random)
            NuevoSubElemento = ET.SubElement(NuevoElemento, 'ns0:inputExpression') # Nivel 2 inputExpresion
            Id_Random = '_' + str(uuid.uuid4()).upper()
            NuevoSubElemento.set('id', Id_Random)
            NuevoSubElemento.set('typeRef', TipoVariable)
            EntradaTexto = ET.Element('ns0:text') # Nivel 3 text (nombre concreto de variable)
            EntradaTexto.text = NombreVariable
            NuevoSubElemento.append(EntradaTexto)

            NivelTabla.append(NuevoElemento)


        # Funcion para agregar nuevos Output
        def AgregarOutput(RutaAcceso, NombreVariable, TipoVariable):
                
            NivelTabla = myroot.find(RutaAcceso)
            NuevoElemento = ET.Element('ns0:output') # Nivel unico output
            Id_Random = '_' + str(uuid.uuid4()).upper()
            NuevoElemento.set('id', Id_Random)
            NuevoElemento.set('name', NombreVariable)
            NuevoElemento.set('typeRef', TipoVariable)

            NivelTabla.append(NuevoElemento)


        # Funcion para agregar nuevas Rules
        def AgregarRule(RutaAcceso, ListaContenido, CantInput, CantOutput):
                
            NivelTabla = myroot.find(RutaAcceso)
            NuevoElemento = ET.Element('ns0:rule') # Nivel 1 rule
            Id_Random = '_' + str(uuid.uuid4()).upper()
            NuevoElemento.set('id', Id_Random)
            
            for i in range(CantInput):
                NuevoSubElemento = ET.SubElement(NuevoElemento, 'ns0:inputEntry') # Nivel 2 inputEntry
                Id_Random = '_' + str(uuid.uuid4()).upper()
                NuevoSubElemento.set('id', Id_Random)
                EntradaTexto = ET.Element('ns0:text') # Nivel 3 text contenido celda regla
                EntradaTexto.text = ListaContenido[i]
                NuevoSubElemento.append(EntradaTexto)

            for i in range(CantOutput):
                NuevoSubElemento = ET.SubElement(NuevoElemento, 'ns0:outputEntry') # Nivel 2 outputEntry
                Id_Random = '_' + str(uuid.uuid4()).upper()
                NuevoSubElemento.set('id', Id_Random)
                EntradaTexto = ET.Element('ns0:text') # Nivel 3 text contenido celda regla
                EntradaTexto.text = ListaContenido[CantInput + i]
                NuevoSubElemento.append(EntradaTexto)

            NivelTabla.append(NuevoElemento)


        global ListaTablasDecision

        NombreArchivoDMN = self.TextoArchivoDMN.text()
        NombreArchivoExcel = self.TextoArchivoExcel.text()
        NombreTabla = self.ListaTablas.currentText()


        self.MensajeError = QMessageBox()
        self.MensajeError.setIcon(QMessageBox.Critical)

        if NombreArchivoExcel == '' or NombreArchivoDMN == '':
            self.MensajeError.setWindowTitle("Error")
            self.MensajeError.setText("Falta seleccionar archivo.")
            self.MensajeError.show()
            
        else:
            ContenidoXLSX = LeerTablaExcel(NombreArchivoExcel)
            if ContenidoXLSX == []:
                self.TextoArchivoExcel.clear()
                self.MensajeError.setWindowTitle("Error")
                self.MensajeError.setText("Excel con formato incorrecto.")
                self.MensajeError.show()
            else:
                mytree = ET.parse(NombreArchivoDMN)
                myroot = mytree.getroot()
                LDdecisionTable = []
                LDatosClaveTotal = []
                TagGenerico = '{http://www.omg.org/spec/DMN/20180521/MODEL/}'
                DatoAgregar = ''

                IndiceSelector = self.ListaTablas.currentIndex()

                SubListaTablas = ListaTablasDecision[IndiceSelector]

                # Leer ruta acceso a tabla desde la lista de datos tablas
                RutaAcceso = SubListaTablas[2]

                #RutaAcceso = './' + TagGenerico + "decision[@name='" + NombreTabla + "']/" + TagGenerico + 'decisionTable'
                
                EliminarSeccion(RutaAcceso, TagGenerico, 'input')
                EliminarSeccion(RutaAcceso, TagGenerico, 'output')
                EliminarSeccion(RutaAcceso, TagGenerico, 'rule')
                EliminarSeccion(RutaAcceso, TagGenerico, 'annotation')

                # Agregar nuevos Input
                for SubListaInput in ContenidoXLSX[0]:
                    AgregarInput(RutaAcceso, SubListaInput[2], SubListaInput[1])

                # Agregar nuevos Output
                for SubListaOutput in ContenidoXLSX[1]:
                    AgregarOutput(RutaAcceso, SubListaOutput[2], SubListaOutput[1])

                # Agregar nuevas rule
                for SubListaRules in ContenidoXLSX[2]:
                    AgregarRule(RutaAcceso, SubListaRules, len(ContenidoXLSX[0]), len(ContenidoXLSX[1]))

                # Guardar resultado final
                mytree.write('resultado.dmn')

                # Reabrir con LXML y mejorar formato
                parser = ETL.XMLParser(remove_blank_text=True)
                NuevoXML = ETL.parse('resultado.dmn', parser)
                NuevoXML.write('resultado.dmn', pretty_print=True)

                QMessageBox.about(self, "OK", "Tabla reemplazada.")

    def ExportarExcel(self):

        fileName = self.TextoArchivoDMN.text()
        nombreTabla = self.ListaTablas.currentText()

        mytree = ET.parse(fileName)
        myroot = mytree.getroot()

        TagGenerico = '{http://www.omg.org/spec/DMN/20180521/MODEL/}'

        global ListaTablasDecision

        IndiceSelector = self.ListaTablas.currentIndex()

        # Nivel tabla de decision
        SubListaTablas = ListaTablasDecision[IndiceSelector]

        RutaAcceso = SubListaTablas[2]

        #RutaAcceso = './' + TagGenerico + "decision" + "[@name='" + nombreTabla + "']/" + TagGenerico  + "decisionTable/"

        # Leer todos los input
        RutaAccesoInput = RutaAcceso + TagGenerico + "input/" + TagGenerico  + "inputExpression/" + TagGenerico + "text"

        ListaInputs = []

        for Input in myroot.findall(RutaAccesoInput):
            ListaInputs.append(Input.text)


        # Leer todos los output
        RutaAccesoOutput = RutaAcceso + TagGenerico + "output"

        ListaOutputs = []

        for Output in myroot.findall(RutaAccesoOutput):
            try:
                ListaOutputs.append(Output.attrib['name'])
            except:
                ListaOutputs.append(SubListaTablas[4]) # Si el output no tiene name 
                                                       # se asume nombre de variable.
                                                       # Esto sucede con decisionTable
                                                       # que tienen como unica variable
                                                       # de salida (la variable de la entrada 
                                                       # del context). 

        # Leer todas las rule
        RutaAccesoRules = RutaAcceso + TagGenerico + "rule"

        ListaRules = []

        for Rule in myroot.findall(RutaAccesoRules): # nivel rule
            ListaRuleParcial = []
            for Rule2 in Rule: # nivel inputEntry u outputEntry o annotationEntry
                for Rule3 in Rule2: # nivel text
                    ListaRuleParcial.append(Rule3.text)
                
            ListaRules.append(ListaRuleParcial)


        ListaEncabezados = ListaInputs + ListaOutputs

        ListaEncabezados.append('annotation-1')

        CantidadInputs = len(ListaInputs)

        CantidadOutputs = len(ListaOutputs)

        # Exportar resultados a archivo Excel

        ArchivoXLSX = openpyxl.Workbook()

        sheet = ArchivoXLSX.active

        sheet.title = 'RESULTADO'

        sheet.append(ListaEncabezados) # Agregar encabezados

        for LineaRule in ListaRules: # Agregar cada linea de rule
            sheet.append(LineaRule)

        # Agregar bordes
        thin_border = Border(
            left=Side(border_style='thin', color='00000000'),
            right=Side(border_style='thin', color='00000000'),
            top=Side(border_style='thin', color='00000000'),
            bottom=Side(border_style='thin', color='00000000')
        )

        for Linea in sheet.iter_rows():
            NumeroLinea = Linea[0].row
            for Celda in Linea:
                NumeroCelda = Celda.column
                Celda.border = thin_border # Agregar borde a cada celda
                if NumeroLinea == 1: # Ajustar texto solo de la linea 1 (encabezados)
                    Celda.alignment = Alignment(wrap_text=True,vertical='center')
                    Celda.font = Font(bold=True)
                    # Aplicar colores a celdas primera linea
                    if NumeroCelda <= CantidadInputs:
                        Celda.fill = PatternFill("solid", start_color="def3ff")
                    elif NumeroCelda > CantidadInputs:
                        Celda.fill = PatternFill("solid", start_color="bee1f4")

        ArchivoXLSX.save('resultado.xlsx')

        QMessageBox.about(self, "OK", "Exportado a archivo excel.")


    def ExplicarDMN(self):

        # Leer archivo DMN y determinar lista de tablas de decision

        fileName = self.TextoArchivoDMN.text()

        NombreArchivoLimp = os.path.basename(fileName)
        NombreArchivoLimp = NombreArchivoLimp.replace('.dmn', '')

        mytree = ET.parse(fileName)
        myroot = mytree.getroot()

        TagGenerico = '{http://www.omg.org/spec/DMN/20180521/MODEL/}'

        global ListaTablasDecision

        if ListaTablasDecision == []:
            print("DMN incorrecto o sin tabla de decisión.")

        mydoc = docx.Document()

        styles = mydoc.styles

        style = styles.add_style('Parrafo', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        style = mydoc.styles['List Bullet 2']
        font = style.font
        font.name = 'Calibri'

        style = mydoc.styles['Heading 2']
        font = style.font
        font.underline = True


        # Agregar encabezado y logo
        header = mydoc.sections[0].header
        paragraph = header.paragraphs[0]

        paragraph.style.font.name = 'Calibri'
        paragraph.text = 'SanCor Salud - DMN\t\t'

        logo_run = paragraph.add_run()
        logo_run.add_picture('logo.png', width=Inches(1))

        header.add_paragraph('\n')

        # Titulo y explicacion general DMN

        mydoc.add_heading('DMN ' + NombreArchivoLimp, 0)

        if len(ListaTablasDecision) > 1:
            mydoc.add_paragraph('El DMN ' + NombreArchivoLimp + ' contiene las siguientes tablas de decision: ' + '\n', style='Parrafo')
        elif len(ListaTablasDecision) == 1:
            mydoc.add_paragraph('El DMN ' + NombreArchivoLimp + ' contiene la siguiente tabla de decision: ' + '\n', style='Parrafo')

        for Tabla in ListaTablasDecision:
            mydoc.add_paragraph(Tabla[0], style='List Bullet 2')



        # Leer y traducir el contenido de cada tabla
        for ITabla, Tabla in enumerate(ListaTablasDecision):

            mydoc.add_page_break() # Agregar salto de pagina desde 2da tabla

            nombreTabla = Tabla[0]

            TagGenerico = '{http://www.omg.org/spec/DMN/20180521/MODEL/}'

            # Nivel tabla de decision
            #RutaAcceso = './' + TagGenerico + "decision" + "[@name='" + nombreTabla + "']/" + TagGenerico  + "decisionTable/"
            
            RutaAcceso = Tabla[2]

            # Leer todos los input
            RutaAccesoInput = RutaAcceso + '/' + TagGenerico + "input/" + TagGenerico  + "inputExpression/" + TagGenerico + "text"

            ListaInputs = []

            for Input in myroot.findall(RutaAccesoInput):
                ListaInputs.append(Input.text)


            # Leer todos los output
            RutaAccesoOutput = RutaAcceso + '/' + TagGenerico + "output"

            ListaOutputs = []

            for Output in myroot.findall(RutaAccesoOutput):
                try:
                    ListaOutputs.append(Output.attrib['name'])
                except:
                    ListaOutputs.append(Tabla[4])


            # Leer todas las rule
            RutaAccesoRules = RutaAcceso + '/' + TagGenerico + "rule"

            ListaRules = []

            for Rule in myroot.findall(RutaAccesoRules): # nivel rule
                ListaRuleParcial = []
                for Rule2 in Rule: # nivel inputEntry u outputEntry o annotationEntry
                    for Rule3 in Rule2: # nivel text
                        ListaRuleParcial.append(Rule3.text)
                    
                ListaRules.append(ListaRuleParcial)


            ListaEncabezados = ListaInputs + ListaOutputs

            ListaEncabezados.append('annotation-1')

            CantidadInputs = len(ListaInputs)

            CantidadOutputs = len(ListaOutputs)



            # Explicar cada rule de la tabla de decision

            # Leer lista de sinonimos
            ListaSin = []
            DiccioSin = {}

            with open('listas/SINONIMOS.txt','r', encoding='ISO-8859-1') as lista:
                for parte in lista.read().splitlines():
                    ListaSin.append(list(parte.split(',')))

                for indice, parte in enumerate(ListaSin):
                    DiccioSin[ListaSin[indice][0]] = str(ListaSin[indice][1]) # Crear diccionario de sinonimos


            # Leer lista de agregados

            ListaAgre = []

            with open('listas/AGREGADOS.txt','r', encoding='ISO-8859-1') as lista:
                for parte in lista.read().splitlines():
                    ListaAgre.append(list(parte.split(',')))



            ExplicacionGeneral = ''

            ExplicacionGeneral = 'Tabla de decision: ' + nombreTabla + '\n\n'


            def TraducirLogica(Contenido): # Funcion para traducir logica basica FEEL

                wordDic = {
                '!=': ' DISTINTO a ',
                'not': ' DISTINTO a ',
                '=<': ' es MENOR o IGUAL que ',
                '>=': ' es MAYOR o IGUAL que ',
                '<': ' es MENOR que ',
                '>': ' es MAYOR que ',
                '+': ' MÁS '}

                ContenidoLimpio = Contenido.replace('  ', ' ')
                ContenidoLimpio = ContenidoLimpio.replace('"', '')
                ContenidoLimpio = ContenidoLimpio.replace("'", '')

                AplicaDiccionario = False

                for key in wordDic:
                    if key in ContenidoLimpio:
                        ContenidoLimpio = ContenidoLimpio.replace(key, wordDic[key])
                        AplicaDiccionario = True
                
                ContenidoLimpio = ContenidoLimpio.replace('(', '')
                ContenidoLimpio = ContenidoLimpio.replace(')', '')

                if AplicaDiccionario == True:
                    TextoTraducido = ContenidoLimpio
                else:
                    TextoTraducido = ' es ' + ContenidoLimpio
                
                return TextoTraducido


            def ReemplazarPorSinonimos(Contenido):

                for key in DiccioSin:
                    Contenido = Contenido.replace(key, DiccioSin[key])

                return Contenido


            def LimpiezaTexto(Texto):

                Texto = Texto.replace('  ', ' ') # Reemplazar doble espacio
                Texto = Texto.replace(' , ', ', ')

                return Texto



            mydoc.add_heading(nombreTabla, 0) # Agregar docx nombre tabla de decision


            for Indice, LineaRegla in enumerate(ListaRules): # nivel de cada linea de regla
                ExplicacionLinea = ''
                TituloExplicativo = str(LineaRegla[-1]) # Tomar annotation como titulo explicativo
                if TituloExplicativo in ('-', '', ' ', 'None'):
                    TituloExplicativo = ''
                TituloLinea = '\n' + 'Regla n° ' + str(Indice) + ': ' + TituloExplicativo + '\n'
                mydoc.add_heading(TituloLinea, 2)
                PrimeraCondicionEncontrada = False
                PrimerResultadoEncontrado = False
                for Indice2, Dato in enumerate(LineaRegla): # nivel de cada dato de la linea
                    if Dato not in ('-', '', ' ', None): # Omitir datos que no suman a la regla
                        for IAgre, Agregado in enumerate(ListaAgre): # Agregados a los textos
                            if ListaEncabezados[Indice2] == Agregado[0]:
                                if Agregado[1] == Dato:
                                    Dato = Dato + ' ' + Agregado[2]
                        if Indice2 in range(CantidadInputs): # Bloque de los inputs o condiciones de la regla
                            if PrimeraCondicionEncontrada == False:
                                ExplicacionLinea = 'Si ' + ListaEncabezados[Indice2] + TraducirLogica(str(Dato))
                                PrimeraCondicionEncontrada = True
                            elif PrimeraCondicionEncontrada == True:
                                ExplicacionLinea = ExplicacionLinea + ', ' + ListaEncabezados[Indice2] + TraducirLogica(str(Dato))
                        elif Indice2 in range((CantidadInputs),(CantidadInputs+CantidadOutputs)): # Bloque de los outputs o resultados de la regla
                            if PrimerResultadoEncontrado == False:
                        
                                ExplicacionLinea = ReemplazarPorSinonimos(ExplicacionLinea) #  Reemplazar por sinonimos
                                ExplicacionLinea = ExplicacionLinea + ', la regla devuelve los siguientes resultados:'
                                ExplicacionLinea = LimpiezaTexto(ExplicacionLinea)
                                mydoc.add_paragraph(ExplicacionLinea, style='Parrafo')
                                mydoc.add_paragraph(ListaEncabezados[Indice2] + ': ' + str(Dato), style='List Bullet 2')

                                ExplicacionLinea = ExplicacionLinea + '* ' + ListaEncabezados[Indice2] + ': ' + str(Dato) + '\n'
                                PrimerResultadoEncontrado = True
                            elif PrimerResultadoEncontrado == True:
                                mydoc.add_paragraph(ListaEncabezados[Indice2] + ': ' + str(Dato), style='List Bullet 2')
                                ExplicacionLinea = ExplicacionLinea + '* ' + ListaEncabezados[Indice2] + ': ' + str(Dato) + '\n'


                
                ExplicacionGeneral = ExplicacionGeneral + '\n\n' + TituloLinea + '\n' + ExplicacionLinea + '\n'

                

                #mydoc.add_paragraph(ExplicacionLinea, style='Parrafo')

                
            ExplicacionGeneral = LimpiezaTexto(ExplicacionGeneral)

        # with open('Explicacion.txt', 'w', encoding='utf-16') as archivo:
        #     archivo.write(ExplicacionGeneral)
        #     archivo.close()


        mydoc.save('resultado.docx')

        QMessageBox.about(self, "OK", "Explicación generada.")


#--------------------Inicio del programa--------------------

app = QtWidgets.QApplication(sys.argv) # Create an instance of QtWidgets.QApplication
window = Ui() # Create an instance of our class
app.exec_() # Start the application


